# -*-coding=utf-8-*-
__author__ = 'milhaven1733'
# -*-coding=utf-8-*-
from email.mime.text import MIMEText
from email.mime.multipart import  MIMEMultipart
from bs4 import BeautifulSoup
from email.Header import Header
from docx import Document
import smtplib
from email import Encoders,Utils
import urllib2
import time
import re
import sys
import os

reload(sys)
sys.setdefaultencoding('utf-8')

def send_to_kindle(smtp_server, from_mail, password, to_mail,filename):
    username=from_mail.split("@")[0]
    #连接邮箱服务
    smtp=smtplib.SMTP()
    smtp.connect(smtp_server)
    smtp.login(username,password)
    #构造带附件的待发送邮件
    msg=MIMEMultipart()
    msg['to']=to_mail
    msg['from']=from_mail
    msg['Subject']="Convert"
    msg['Date']=Utils.formatdate(localtime=1)
    filename=filename+'.docx'
    #构造及链接附件
    content=open(filename.decode('utf-8'),'rb').read()
    att=MIMEText(content,'base64','utf-8')#指定邮件内容，格式及字符编码
    att["Content-Disposition"] = 'attachment;filename="%s"'%Header(filename,'gb2312')
    msg.attach(att)
    smtp.sendmail(msg['from'],msg['to'],msg.as_string())
    smtp.quit()

def Getcontent(id):
    host = "http://www.zhihu.com"
    url=host+'/question/'+id
    print url
    user_agent='Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/49.0.2623.75 Safari/537.36'
    header = {"User-Agent": user_agent}
    req=urllib2.Request(url,headers=header)
    try:
        resp=urllib2.urlopen(req)
    except:
        print "Time out. Retry"
        time.sleep(30)
        resp=urllib2.urlopen(req)
       
    try:
        bs=BeautifulSoup(resp,"lxml")
    except:
        print "Beautifulsoup error"
        return None
    
    title=bs.title
    print title
    
    document=Document()
    filename=re.sub('[\/:*?"<>|]', '-',title.string.strip())
    print filename
    
    document.add_paragraph(title.string)
    
    detail=bs.find('div',class_='zm-editable-content')
    if detail is not None:
        document.add_paragraph("\n\n--------------------Detail----------------------\n\n")
        for details in detail.strings:
            document.add_paragraph(unicode(details))
    
    answer = bs.find_all('div',class_="zm-editable-content clearfix")
    k=0
    count=1 
    index=1
    for each_answer in answer:
        image=each_answer.find_all('img',class_="origin_image zh-lightbox-thumb")
        for i in image:
            lips=i.get('src')
            result=urllib2.urlopen(lips)
            data=result.read()
            name=os.getcwd()+'/'+str(count)
            with open(name,'wb')as code:
                code.write(data)
                code.close()
            count+=1
    
        document.add_paragraph("\n\n-------------------------answer %s -------------------------\n\n" % k)
        for answers in each_answer.strings:
            document.add_paragraph(unicode(answers))
        for i in range(index,count):
            document.add_picture(str(i))
        index=count
        k+=1
    document.save(filename+'.docx')
    return filename

    
if __name__=="__main__":
    folder=os.path.join(os.getcwd(), "kindle")
    if not os.path.exists(folder):
        os.mkdir(folder)
        
    os.chdir(folder)
    
    id=sys.argv[1]
    filename=Getcontent(id)
    
    smtp_server='smtp.XX.com'
    from_mail='your_username@XX.com'
    password='your_password' 
    to_mail='XX@kindle.cn'
    send_to_kindle(smtp_server, from_mail, password, to_mail,filename)

    print "已发送至kindle."
