# -*-coding=utf-8-*-
__author__ = 'milhaven1733'
# -*-coding=utf-8-*-
from email.mime.text import MIMEText
from email.mime.multipart import  MIMEMultipart
from docx import Document
import smtplib
from email import Encoders,Utils
import urllib2
import time
import re
import sys
import os

from bs4 import BeautifulSoup
from email.Header import Header

reload(sys)
sys.setdefaultencoding('utf-8')

class Getcontent():
    def __init__(self, id):
        id_link='/question/'+id
        self.getAnswer(id_link)
        
    def getAnswer(self,answerID):
        host = "http://www.zhihu.com"
        url=host+answerID
        print url
        user_agent='Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/49.0.2623.75 Safari/537.36'
        header = {"User-Agent": user_agent}
        req=urllib2.Request(url,headers=header)
        
        try:
            resp=urllib2.urlopen(req)
        except:
            print "Time out. Retry"
            time.sleep(30)
            resp=urllib2.urlopen(req)
           
        try:
            bs=BeautifulSoup(resp,"lxml")
            
        except:
            print "Beautifulsoup error"
            return None

        title=bs.title
        print title
	
	document=Document()
        
        filename_old=title.string.strip()
        print filename_old
        filename = re.sub('[\/:*?"<>|]', '-', filename_old)

	document.add_paragraph(title.string)
        
        detail=bs.find('div',class_='zm-editable-content')
        
	document.add_paragraph("\n\n\n\n--------------------Detail----------------------\n\n")        
        if detail is not None:
            for details in detail.strings:
             	document.add_paragraph(unicode(detail))
   
        answer = bs.find_all('div',class_="zm-editable-content clearfix")
        k=0
	count=1 
	index=1
        for each_answer in answer:
	    image=each_answer.find_all('img',class_="origin_image zh-lightbox-thumb")
	    for i in image:
		lips=i.get('src')
		result=urllib2.urlopen(lips)
		data=result.read()
		name=os.getcwd()+'/'+str(count)
		with open(name,'wb')as code:
		    code.write(data)
		    code.close()
		count+=1
	    document.add_paragraph("\n\n-------------------------answer %s via  -------------------------\n\n" % k)
            for answers in each_answer.strings:
	    	document.add_paragraph(unicode(answers)
	    for i in range(index,count):
		document.add_picture(str(i))
	    index=count
            k+=1
	document.save(filename+'.docx')
	
if __name__=="__main__":
    folder=os.path.join(os.getcwd(), "content")
    if not os.path.exists(folder):
        os.mkdir(folder)
        
    os.chdir(folder)
    
    id=sys.argv[1]

    obj=Getcontent(id)
    print "Done"
